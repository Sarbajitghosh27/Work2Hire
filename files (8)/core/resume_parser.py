"""
core/resume_parser.py  — v3 (anti-hallucination + text normalization)
───────────────────────────────────────────────────────────────────────
Upgrade 4: Parser Cleanup & Text Normalization
  - Fixes CamelCase concatenation artefacts (LibrariesPandas → Libraries Pandas)
  - Normalises bullet markers, whitespace, OCR garbage
  - Adds sparse_sections detection for Streamlit onboarding hints

Contact info extracted from HEADER ONLY.
Every other section is sanitized to remove any leaking contact strings.
"""

import re
import io
import logging
from dataclasses import dataclass, field
from typing import List, Optional

import pdfplumber
from docx import Document

from config import canonical_skill_name

logger = logging.getLogger(__name__)

EMAIL_RE    = re.compile(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}", re.I)

# ──────────────────────────────────────────────────────────────────────────────
# TEXT NORMALISATION (Upgrade 4)
# ──────────────────────────────────────────────────────────────────────────────
_OCR_GARBAGE_RE = re.compile(r"[^\x20-\x7E\n\t]")   # non-printable ASCII
_MULTI_SPACE_RE = re.compile(r"[ \t]{2,}")
_MULTI_NL_RE    = re.compile(r"\n{3,}")
# Fix PDF CamelCase artefacts: LibrariesPandas → Libraries Pandas
_CAMEL_SPLIT_RE = re.compile(r"([a-z])([A-Z])")


def normalize_text(text: str) -> str:
    """
    Normalise raw extracted text before parsing.
    Fixes common PDF/OCR artefacts.
    """
    # Fix CamelCase concatenations from PDF extraction (LibrariesPandas → Libraries Pandas)
    text = _CAMEL_SPLIT_RE.sub(r"\1 \2", text)

    # Normalize unicode bullets / dashes to ASCII
    text = text.replace("\u2022", "-").replace("\u25cf", "-").replace("\u2013", "-").replace("\u2014", " - ")

    # Strip non-printable characters (OCR garbage)
    text = _OCR_GARBAGE_RE.sub(" ", text)

    # Normalize whitespace
    text = _MULTI_SPACE_RE.sub(" ", text)
    text = _MULTI_NL_RE.sub("\n\n", text)

    # Normalize bullet markers to consistent form
    text = re.sub(r"^[ \t]*[>*•●▪◦◆■□]\s+", "- ", text, flags=re.M)

    return text.strip()
PHONE_RE    = re.compile(r"(\+?\d[\d\s\-().]{7,15}\d)")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.I)
GITHUB_RE   = re.compile(r"github\.com/[\w\-]+", re.I)
URL_RE      = re.compile(r"https?://\S+", re.I)
METRIC_RE   = re.compile(r"\d+[\.,]?\d*\s*(%|x|X|\+|k|K|M|B|ms|sec|hrs?|users?|requests?)", re.I)
BULLET_RE   = re.compile(r"^[\s]*[•\-\*\u2022\u25cf>]\s+(.+)$", re.M)
_MONTH_NAME  = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*"
_SINGLE_DATE = rf"(?:{_MONTH_NAME}[\s,]*\d{{2,4}}|\d{{1,2}}/\d{{2,4}}|\d{{1,2}}-\d{{2,4}}|\d{{4}})"
_DATE_SEP    = r"\s*(?:[-–—]|to)\s*"
_END_DATE    = rf"(?:{_SINGLE_DATE}|present|current|now)"
DATE_RE      = re.compile(rf"\b{_SINGLE_DATE}(?:{_DATE_SEP}{_END_DATE})?\b", re.I)


SECTION_PATTERNS = {
    "summary"        : re.compile(r"^\s*(summary|objective|profile|about me|overview)\s*$", re.I | re.M),
    "experience"     : re.compile(r"^\s*(work experience|experience|employment|professional experience|internship experience|volunteering experience)\s*$", re.I | re.M),
    "projects"       : re.compile(r"^\s*(projects?|portfolio|personal projects?)\s*$", re.I | re.M),
    "education"      : re.compile(r"^\s*(education|academic|qualifications?)\s*$", re.I | re.M),
    "skills"         : re.compile(r"^\s*(skills?|technical skills?|technologies|competencies|tech stack)\s*$", re.I | re.M),
    "certifications" : re.compile(r"^\s*(certifications?|certificates?|credentials?|awards?)\s*$", re.I | re.M),
    "accomplishments": re.compile(r"^\s*(accomplishments?|achievements?|extra.?curricular|activities|volunteering)\s*$", re.I | re.M),
    "publications"   : re.compile(r"^\s*(publications?|research|papers?)\s*$", re.I | re.M),
}

DEGREE_KW = ["b.tech","m.tech","b.e","m.e","bsc","msc","mba","phd",
             "bachelor","master","diploma","b.sc","m.sc","b.com","mca","be","me"]
DEGREE_RE = re.compile(r'\b(?:' + '|'.join(re.escape(kw) for kw in DEGREE_KW) + r')\b', re.I)


@dataclass
class WorkExperience:
    company  : str = ""
    title    : str = ""
    duration : str = ""
    bullets  : List[str] = field(default_factory=list)

@dataclass
class Project:
    name       : str = ""
    tech_used  : List[str] = field(default_factory=list)
    description: str = ""
    outcome    : str = ""
    link       : str = ""
    duration   : str = ""


@dataclass
class Education:
    institution: str = ""
    degree     : str = ""
    year       : str = ""
    gpa        : str = ""

@dataclass
class ParsedResume:
    raw_text       : str = ""
    name           : str = ""
    email          : str = ""
    phone          : str = ""
    linkedin       : str = ""
    github         : str = ""
    portfolio      : str = ""
    summary        : str = ""
    skills         : List[str] = field(default_factory=list)
    experience     : List[WorkExperience] = field(default_factory=list)
    projects       : List[Project] = field(default_factory=list)
    education      : List[Education] = field(default_factory=list)
    certifications : List[str] = field(default_factory=list)
    accomplishments: List[str] = field(default_factory=list)
    publications   : List[str] = field(default_factory=list)
    all_keywords   : List[str] = field(default_factory=list)
    metrics_found  : List[str] = field(default_factory=list)
    word_count     : int = 0
    sections_found : List[str] = field(default_factory=list)
    sparse_sections: List[str] = field(default_factory=list)  # Upgrade 4: thin-section detection


def extract_text_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        logger.error(f"PDF fail: {e}")
    return text.strip()


def extract_text_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)


def extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.lower().endswith(".pdf"):
        return extract_text_pdf(file_bytes)
    elif filename.lower().endswith((".docx", ".doc")):
        return extract_text_docx(file_bytes)
    raise ValueError(f"Unsupported: {filename}")


def split_into_sections(text: str) -> dict:
    lines = text.split("\n")
    section_starts = {}
    for i, line in enumerate(lines):
        for sec_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(line.strip()) and sec_name not in section_starts.values():
                section_starts[i] = sec_name
                break

    if not section_starts:
        return {"header": text}

    sorted_starts = sorted(section_starts.items())
    first_line    = sorted_starts[0][0]
    result        = {"header": "\n".join(lines[:first_line]).strip()}

    for idx, (line_no, sec_name) in enumerate(sorted_starts):
        end = sorted_starts[idx + 1][0] if idx + 1 < len(sorted_starts) else len(lines)
        result[sec_name] = "\n".join(lines[line_no + 1: end]).strip()

    return result


def extract_contact_from_header(header: str) -> dict:
    c = {"name": "", "email": "", "phone": "", "linkedin": "", "github": "", "portfolio": ""}
    m = EMAIL_RE.search(header);    c["email"]    = m.group()                       if m else ""
    m = LINKEDIN_RE.search(header); c["linkedin"] = "https://" + m.group()          if m else ""
    m = GITHUB_RE.search(header);   c["github"]   = "https://" + m.group()          if m else ""
    
    # Portfolio: find any URL in the header that is not LinkedIn or GitHub
    for m in URL_RE.finditer(header):
        url = m.group()
        if "linkedin.com" not in url.lower() and "github.com" not in url.lower():
            c["portfolio"] = url
            break

    # Phone: must have >=7 digits
    for m in PHONE_RE.finditer(header):
        digits = re.sub(r"\D", "", m.group())
        if len(digits) >= 7:
            c["phone"] = m.group().strip()
            break
    # Name: first short title-case line with no special chars
    for line in header.split("\n")[:10]:
        line = line.strip()
        words = line.split()
        if (2 <= len(words) <= 5
                and not any(ch in line for ch in ["@","http","+","|","/","\\"])
                and not any(ch.isdigit() for ch in line)):
            c["name"] = line
            break
    return c


def build_contact_pattern(c: dict) -> Optional[re.Pattern]:
    frags = []
    if c["email"]:    frags.append(re.escape(c["email"]))
    if c["phone"]:
        digits = re.sub(r"\D", "", c["phone"])
        if len(digits) >= 7: frags.append(re.escape(digits[:8]))
    if c["linkedin"]: frags.append(re.escape(c["linkedin"].replace("https://", "")))
    if c["github"]:   frags.append(re.escape(c["github"].replace("https://", "")))
    if c.get("portfolio"): frags.append(re.escape(c["portfolio"].replace("https://", "").replace("http://", "")))
    return re.compile("|".join(frags), re.I) if frags else None


def sanitize(text: str, pat: Optional[re.Pattern]) -> str:
    if not pat:
        return text
    return "\n".join(
        line for line in text.split("\n")
        if not pat.search(line)
    )


def extract_skills(skills_text: str, full_text: str) -> List[str]:
    from config import TECH_TAXONOMY
    all_tech = [t for terms in TECH_TAXONOMY.values() for t in terms]
    combined = (skills_text + "\n" + full_text).lower()
    found = []
    for tech in all_tech:
        if re.search(r"\b" + re.escape(tech) + r"\b", combined):
            found.append(tech)

    BANNED_SKILL_WORDS = {
        "cross-functional", "collaboration", "project management", "problem solving", "nmos", "gain",
        "machine learning", "embedded systems", "bit mesra", "mesra", "1st position", "first position",
        "leadership", "engineering projects", "conference presentations", "experience & leadership",
        "awards & certifications", "awards", "certifications", "online courses", "competitions",
        "scholarships", "academic or professional recognitions", "used python", "data cleaning",
        "transparency", "refining", "alerts", "gps", "basic", "studio", "1st", "position",
        "custom vlsi layout", "hardware programming", "vlsi layout", "hardware programming"
    }

    # Also parse comma/bullet list from skills section
    for sep in [",", "|", "•", "/", "\n"]:
        for part in skills_text.split(sep):
            clean = part.strip().strip("•-*.,()[]{} ").lower()
            # Clean leading transition words
            clean = re.sub(r"^(and|or|with|in|to)\s+", "", clean).strip()
            clean = clean.strip("•-*.,()[]{} ")
            
            if not clean or len(clean) < 2 or len(clean) > 35 or len(clean.split()) > 4:
                continue

            is_banned = False
            for banned in BANNED_SKILL_WORDS:
                if clean == banned or (len(banned) > 3 and banned in clean):
                    is_banned = True
                    break

            if is_banned:
                continue

            if clean not in found:
                found.append(clean)

    # Deduplicate case-insensitively: keep first occurrence, prefer title-case
    seen_lower: set = set()
    deduped = []
    for s in found:
        s_clean = re.sub(r"^(and|or|with|in|to)\s+", "", s.strip().strip("•-*.,()[]{} ")).strip()
        key = s_clean.lower().strip()
        if not key or key in seen_lower:
            continue
            
        is_banned = False
        for banned in BANNED_SKILL_WORDS:
            if key == banned or (len(banned) > 3 and banned in key):
                is_banned = True
                break

        if is_banned:
            continue

        seen_lower.add(key)
        deduped.append(canonical_skill_name(s_clean))
    return deduped[:50]


# Placeholder strings that indicate the parser failed — must be rejected
_PLACEHOLDER_PATTERNS = re.compile(
    r"(recent employer|company name|job title|professional role|\bN/A\b|"
    r"[Rr]ecent [Ee]mployer|[Cc]ompany [Nn]ame|[Jj]ob [Tt]itle)",
    re.I,
)

def _is_contact_line(line: str, contact_pat: Optional[re.Pattern]) -> bool:
    """True if a line is purely contact info — should not be a job title."""
    if contact_pat and contact_pat.search(line):
        return True
    # Pure digits / phone-number line
    if re.match(r"^[\d\s\-().+]+$", line.strip()) and len(re.sub(r"\D", "", line)) >= 7:
        return True
    if re.match(r"^(https?://|linkedin|github|@)", line.strip(), re.I):
        return True
    return False


def extract_experience(exp_text: str, contact_pat: Optional[re.Pattern] = None) -> List[WorkExperience]:
    entries = []
    blocks  = re.split(r"\n{2,}", exp_text.strip())
    for block in blocks:
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        if not lines or len(block) < 10:
            continue
        we = WorkExperience()

        # 1. Extract duration using DATE_RE
        duration = ""
        for line in lines:
            dm = DATE_RE.search(line)
            if dm:
                duration = dm.group().strip()
                break
        we.duration = duration

        # 2. Clean the date range out of header lines
        cleaned_lines = list(lines)
        for i in range(min(2, len(lines))):
            cleaned = DATE_RE.sub("", lines[i])
            cleaned = re.sub(r"\s*[-–—|]\s*$", "", cleaned)
            cleaned = re.sub(r"^\s*[-–—|]\s*", "", cleaned)
            cleaned_lines[i] = cleaned.strip()

        header = cleaned_lines[0]

        # 3. Reject blocks where the header is contact info or a placeholder
        if _is_contact_line(header, contact_pat) or _PLACEHOLDER_PATTERNS.search(header):
            continue
        # Also skip if header is empty after cleaning
        if not header or len(header) < 3:
            continue

        if " | " in header or " – " in header or " - " in header:
            parts = re.split(r" \| | – | - ", header, maxsplit=1)
            we.title   = parts[0].strip()
            we.company = parts[1].strip() if len(parts) > 1 else ""
        elif " at " in header.lower():
            parts = re.split(r" at ", header, flags=re.I, maxsplit=1)
            we.title   = parts[0].strip()
            we.company = parts[1].strip()
        else:
            # header is company, next line might be title
            candidate_title = cleaned_lines[1] if len(cleaned_lines) > 1 else ""
            # Don't use it as a title if it looks like a bullet or contact
            if candidate_title and not re.match(r"^[\-\*•]", candidate_title) \
                    and not _is_contact_line(candidate_title, contact_pat) \
                    and not _PLACEHOLDER_PATTERNS.search(candidate_title):
                we.company = header
                we.title   = candidate_title
            else:
                we.company = header
                we.title   = ""

        # 4. Reject placeholder titles/companies
        if _PLACEHOLDER_PATTERNS.search(we.title) or _PLACEHOLDER_PATTERNS.search(we.company):
            continue

        bullets = BULLET_RE.findall(block)
        we.bullets = [b.strip() for b in bullets if len(b.strip()) > 10] or \
                     [l for l in lines[1:] if l and not DATE_RE.search(l)
                      and len(l) > 10
                      and not _is_contact_line(l, contact_pat)][:8]
        if we.company or we.title:
            entries.append(we)
    return entries


def _is_project_title(line: str) -> bool:
    """
    True if this line looks like a genuine project title.
    A title must NOT be:
      - A bullet point
      - A tech-stack line
      - A short metric/continuation fragment
      - Starting with a lowercase word (continuation sentence)
      - Very long (>120 chars — that's a bullet, not a title)
    """
    line = line.strip()
    if not line or len(line) < 8 or len(line) > 120:
        return False
    # Bullet lines
    if re.match(r"^[\-\*\u2022\u25cf>]", line):
        return False
    # Tech stack lines
    if re.match(r"^(Tech Stack|Stack|Tools?|Tech)\s*:", line, re.I):
        return False
    # Continuation fragments — start with lowercase or common connector words
    if re.match(r"^(achieving|using|with|for|and|via|by|to |the |a |an )", line, re.I):
        return False
    first_word = line.split()[0] if line.split() else ""
    if first_word and first_word[0].islower():
        return False
    # Lines that are clearly metric/result fragments
    _CONTINUATION_WORDS = re.compile(
        r"^(accuracy|recall|precision|auc|f1|false positive|latency|"
        r"detection|analysis|quantification|infrastructure|thresholds|"
        r"deployment|pipeline|context.aware|fabrication)",
        re.I,
    )
    if _CONTINUATION_WORDS.match(line):
        return False
    return True


def extract_projects(proj_text: str, skills: List[str]) -> List[Project]:
    """
    Rewritten project extractor — no hallucination.
    
    Strategy: walk lines one at a time.
    A new project starts when we see a TITLE line after having seen bullet content.
    Everything between titles = that project's body.
    Hard limit: 5 projects max (you can't have 10 projects in a 1-page CV).
    """
    projects = []
    lines = [l.rstrip() for l in proj_text.split("\n")]

    current_title = None
    current_body = []
    had_bullets = False
    first_entry = True

    def _flush(title, body):
        """Parse a collected project block and return a Project or None."""
        if not title or len(title.strip()) < 5:
            return None
        proj = Project()
        
        # Extract duration from project title if present
        duration = ""
        dm = DATE_RE.search(title)
        if dm:
            duration = dm.group().strip()
            title = DATE_RE.sub("", title)
            title = re.sub(r"\s*\(\s*\)\s*", " ", title)
            title = re.sub(r"\s*[-–—|]\s*$", "", title)
            title = re.sub(r"^\s*[-–—|]\s*", "", title)
            title = title.strip().rstrip(",:")
        else:
            # Check first line of body for a date
            if body:
                first_body_line = body[0]
                dm = DATE_RE.search(first_body_line)
                if dm:
                    duration = dm.group().strip()
                    body[0] = DATE_RE.sub("", first_body_line).strip()
                    
        proj.name = title
        proj.duration = duration

        # Tech stack: explicit "Tech Stack:" line
        for line in body:
            m = re.match(r"^(?:Tech Stack|Stack|Tools?|Tech)\s*:\s*(.+)", line, re.I)
            if m:
                proj.tech_used = [t.strip() for t in re.split(r"[|,]", m.group(1)) if t.strip()][:8]
                break

        # Fallback tech: infer from skills present in body text
        if not proj.tech_used:
            body_text = " ".join(body).lower()
            proj.tech_used = [s for s in skills if re.search(r"\b" + re.escape(s) + r"\b", body_text)][:6]

        # Bullet points → description + outcome
        bullets = []
        for line in body:
            m = re.match(r"^[\-\*\u2022\u25cf>]\s*(.+)", line)
            if m:
                b = m.group(1).strip()
                if len(b) > 15:
                    bullets.append(b)

        if bullets:
            proj.description = bullets[0]
            metric_bullets = [b for b in bullets if METRIC_RE.search(b)]
            proj.outcome = metric_bullets[0] if metric_bullets else (bullets[-1] if len(bullets) > 1 else "")
        else:
            # Non-bulleted lines
            content = [l for l in body
                       if not re.match(r"^(?:Tech Stack|Stack|Tools?|Tech)\s*:", l, re.I)
                       and len(l.strip()) > 15]
            proj.description = " ".join(content[:2])
            metric_lines = [l for l in content if METRIC_RE.search(l)]
            proj.outcome = metric_lines[0] if metric_lines else ""

        # Link
        link_m = URL_RE.search(" ".join(body))
        if link_m:
            proj.link = link_m.group()

        # Only return if there's actual content
        if proj.name and (proj.description or proj.tech_used):
            return proj
        return None

    for raw_line in lines:
        line = raw_line.strip()

        is_bullet = bool(re.match(r"^[\-\*\u2022\u25cf>]", line))
        is_tech   = bool(re.match(r"^(?:Tech Stack|Stack|Tools?|Tech)\s*:", line, re.I))
        is_empty  = len(line) == 0

        if is_empty:
            continue

        if is_bullet or is_tech:
            had_bullets = True
            if current_title is not None:
                current_body.append(line)
            continue

        # Candidate title check
        if _is_project_title(line) and (had_bullets or first_entry):
            # Flush previous
            if current_title is not None:
                p = _flush(current_title, current_body)
                if p:
                    projects.append(p)
            current_title = line
            current_body = []
            had_bullets = False
            first_entry = False
        else:
            # Continuation line
            if current_title is not None:
                current_body.append(line)
            elif _is_project_title(line):
                # First title before any bullets
                current_title = line
                current_body = []
                first_entry = False

    # Flush last
    if current_title:
        p = _flush(current_title, current_body)
        if p:
            projects.append(p)

    # Hard safety cap: a 1-page resume cannot have more than 5 real projects
    return projects[:5]


def extract_education(edu_text: str) -> List[Education]:
    _GPA_NUM_RE  = re.compile(r"(?:sgpa|cgpa|gpa|percentage)[:\s]+([\d.]+(?:\s*/\s*[\d.]+)?)", re.I)
    _GPA_FRAG_RE = re.compile(r"[|\u2192\u2013\u2014]?\s*(?:sgpa|cgpa|gpa|percentage)[:\s]+[\d.]+(?:\s*/\s*[\d.]+)?", re.I)

    entries = []

    # ── Improved splitting: try double-newline first, fall back to per-line ────
    raw_blocks = re.split(r"\n{2,}", edu_text.strip())

    # If we got only 1 big block, split line-by-line and group greedily
    if len(raw_blocks) <= 1:
        all_lines = [l.strip() for l in edu_text.split("\n") if l.strip()]
        # Each institution/degree is typically 1–3 consecutive lines.
        # Regroup: start a new block whenever we see a new institution keyword
        # OR a degree keyword on a new line that has a date.
        raw_blocks = []
        current_block_lines: List[str] = []
        for ln in all_lines:
            ll = ln.lower()
            is_inst  = bool(re.search(r"(university|institute|college|iit|nit|bits|school|board)", ll, re.I))
            is_deg   = bool(DEGREE_RE.search(ll))
            has_date = bool(DATE_RE.search(ln))
            # Start a new block if this line looks like a new entry header
            if (is_inst or (is_deg and has_date)) and current_block_lines:
                raw_blocks.append("\n".join(current_block_lines))
                current_block_lines = []
            current_block_lines.append(ln)
        if current_block_lines:
            raw_blocks.append("\n".join(current_block_lines))

    for block in raw_blocks:
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        if not lines:
            continue
        edu = Education()

        # 1. Extract the FIRST date range in this block only
        year = ""
        for line in lines:
            dm = DATE_RE.search(line)
            if dm and not year:
                year = dm.group().strip()
                break
        edu.year = year

        # 2. Parse degree and institution from lines
        for line in lines:
            cleaned_line = DATE_RE.sub("", line)
            cleaned_line = re.sub(r"\s*[-–—|]\s*$", "", cleaned_line)
            cleaned_line = re.sub(r"^\s*[-–—|]\s*", "", cleaned_line)
            cleaned_line = cleaned_line.strip()
            
            cleaned_ll = cleaned_line.lower()
            has_deg_kw = bool(DEGREE_RE.search(cleaned_ll))
            has_inst_kw = bool(re.search(r"(university|institute|college|iit|nit|bits|school|board)", cleaned_ll, re.I))
            
            parts = [cleaned_line]
            if has_deg_kw and has_inst_kw:
                # Try to split by common separators including "from" and comma
                split_parts = re.split(r"\s*(?:[-–—|]|\bat\b|\bfrom\b|,\s+)\s*", cleaned_line, flags=re.I)
                if len(split_parts) > 1:
                    parts = [p.strip() for p in split_parts if p.strip()]
            
            for part in parts:
                part_ll = part.lower()
                is_deg = bool(DEGREE_RE.search(part_ll))
                is_inst = bool(re.search(r"(university|institute|college|iit|nit|bits|school|board)", part_ll, re.I))

                if is_deg and not edu.degree:
                    clean_deg = _GPA_FRAG_RE.sub("", part)
                    clean_deg = re.sub(r"\s*\(\s*\)", "", clean_deg)
                    clean_deg = re.sub(r"[|\u2192\u2013\u2014]+", " ", clean_deg)
                    clean_deg = re.sub(r"\s{2,}", " ", clean_deg).strip().rstrip("-–— |")
                    edu.degree = clean_deg

                if is_inst and not edu.institution:
                    clean_inst = _GPA_FRAG_RE.sub("", part)
                    clean_inst = re.sub(r"\s*\(\s*\)", "", clean_inst)
                    clean_inst = re.sub(r"[|\u2192]+", " ", clean_inst)
                    edu.institution = re.sub(r"\s{2,}", " ", clean_inst).strip()

            gm = _GPA_NUM_RE.search(line)
            if gm and not edu.gpa:
                edu.gpa = gm.group(1).strip()

        if edu.degree or edu.institution:
            entries.append(edu)

    # Deduplicate entries with identical or highly similar institution (case-insensitive, alphanumeric check)
    unique: List[Education] = []
    for e in entries:
        found_idx = -1
        edu_deg = (e.degree or "").strip()
        edu_inst = (e.institution or "").strip()
        edu_yr = (e.year or "").strip()
        edu_gpa = (e.gpa or "").strip()
        edu_norm = re.sub(r"[^a-z0-9]", "", (edu_deg + edu_inst).lower())

        for idx, existing in enumerate(unique):
            ex_deg = (existing.degree or "").strip()
            ex_inst = (existing.institution or "").strip()
            ex_norm = re.sub(r"[^a-z0-9]", "", (ex_deg + ex_inst).lower())

            # Check year conflict
            yr1 = re.sub(r"[^0-9]", "", edu_yr)
            yr2 = re.sub(r"[^0-9]", "", (existing.year or ""))
            years_conflict = False
            if yr1 and yr2:
                y1_set = set(re.findall(r"\b\d{4}\b", edu_yr))
                y2_set = set(re.findall(r"\b\d{4}\b", existing.year or ""))
                if y1_set and y2_set and not (y1_set & y2_set):
                    years_conflict = True

            if years_conflict:
                continue

            is_match = False
            if edu_norm and ex_norm:
                if edu_norm == ex_norm or edu_norm in ex_norm or ex_norm in edu_norm:
                    is_match = True

            inst1 = re.sub(r"[^a-z0-9]", "", edu_inst.lower())
            inst2 = re.sub(r"[^a-z0-9]", "", ex_inst.lower())
            if inst1 and inst2 and (inst1 in inst2 or inst2 in inst1):
                is_match = True

            if not inst1 and inst2:
                deg1_norm = re.sub(r"[^a-z0-9]", "", edu_deg.lower())
                if inst2 in deg1_norm:
                    is_match = True
            if not inst2 and inst1:
                deg2_norm = re.sub(r"[^a-z0-9]", "", ex_deg.lower())
                if inst1 in deg2_norm:
                    is_match = True

            if is_match:
                found_idx = idx
                break

        if found_idx != -1:
            existing = unique[found_idx]
            if len(edu_deg) > len(existing.degree or ""):
                existing.degree = edu_deg
            if len(edu_inst) > len(existing.institution or ""):
                existing.institution = edu_inst
            if not existing.year and edu_yr:
                existing.year = e.year
            if not existing.gpa and edu_gpa:
                existing.gpa = e.gpa
        else:
            unique.append(e)
    return unique


def is_valid_certification(c: str) -> bool:
    c_low = c.lower().strip()
    if len(c) > 100 or len(c.split()) > 12:
        return False
    banned_keywords = [
        "does not list",
        "no certifications",
        "no separate certifications",
        "however, it does include",
        "not specify certifications",
        "not list certifications",
        "no online courses",
        "does not specify",
        "no formal certifications",
        "n/a",
        "none",
        "not applicable",
        "available upon request",
        "refer to",
        "see project",
        "no direct certifications",
        "not yet certified",
        "has not listed",
        "does not have",
        "not listed",
        "no separate",
        "no online",
        "will be provided",
        "explanatory text",
        "source document",
        "explicitly mentioned",
        "no direct",
        "no certifications are",
        "refer to ",
        "nil"
    ]
    for kw in banned_keywords:
        if kw in c_low:
            return False
    return True


def is_valid_accomplishment(a: str) -> bool:
    a_low = a.lower().strip()
    if len(a) > 150 or len(a.split()) > 20:
        return False
    banned_keywords = [
        "does not list",
        "no accomplishments",
        "no separate accomplishments",
        "no achievements",
        "no awards",
        "however, it does include",
        "not specify accomplishments",
        "not list accomplishments",
        "does not specify",
        "no formal accomplishments",
        "no formal achievements",
        "n/a",
        "none",
        "not applicable",
        "available upon request",
        "refer to",
        "see project",
        "no direct achievements",
        "no direct accomplishments",
        "has not listed",
        "does not have",
        "not listed",
        "will be provided",
        "explanatory text",
        "source document",
        "explicitly mentioned",
        "nil"
    ]
    for kw in banned_keywords:
        if kw in a_low:
            return False
    return True



def parse_resume(file_bytes: bytes, filename: str) -> ParsedResume:
    parsed = ParsedResume()
    raw    = extract_text(file_bytes, filename)

    # ── Upgrade 4: Normalize raw text before any parsing ─────────────────────
    raw = normalize_text(raw)

    parsed.raw_text   = raw
    parsed.word_count = len(raw.split())

    sections = split_into_sections(raw)
    parsed.sections_found = [k for k in sections if k != "header"]

    # Contact: header ONLY
    header  = sections.get("header", raw[:600])
    contact = extract_contact_from_header(header)
    parsed.name      = contact["name"]
    parsed.email     = contact["email"]
    parsed.phone     = contact["phone"]
    parsed.linkedin  = contact["linkedin"]
    parsed.github    = contact["github"]
    parsed.portfolio = contact["portfolio"]

    cpat = build_contact_pattern(contact)

    parsed.summary        = sanitize(sections.get("summary", ""), cpat).strip()[:600]
    parsed.skills         = extract_skills(sections.get("skills", ""), raw)
    parsed.experience     = extract_experience(sanitize(sections.get("experience", ""), cpat), cpat)
    parsed.projects       = extract_projects(sanitize(sections.get("projects", ""), cpat), parsed.skills)
    parsed.education      = extract_education(sanitize(sections.get("education", ""), cpat))

    cert_raw = sanitize(sections.get("certifications", ""), cpat)
    raw_certs = [b.strip() for b in BULLET_RE.findall(cert_raw) if b.strip()] or \
                [l.strip() for l in cert_raw.split("\n") if l.strip() and len(l) > 5][:10]
    parsed.certifications = [c for c in raw_certs if is_valid_certification(c)]

    acc_raw = sanitize(sections.get("accomplishments", ""), cpat)
    raw_accs = [
        b.strip() for b in BULLET_RE.findall(acc_raw)
        if b.strip() and not EMAIL_RE.search(b) and not LINKEDIN_RE.search(b)
    ] or [
        l.strip() for l in acc_raw.split("\n")
        if l.strip() and len(l.strip()) > 15
        and not EMAIL_RE.search(l) and not LINKEDIN_RE.search(l) and not PHONE_RE.search(l)
    ]
    parsed.accomplishments = [a for a in raw_accs if is_valid_accomplishment(a)]

    pub_raw = sanitize(sections.get("publications", ""), cpat)
    parsed.publications   = [l.strip().rstrip(":,") for l in pub_raw.split("\n") if l.strip() and len(l) > 10][:8]

    parsed.metrics_found  = list(set(METRIC_RE.findall(raw)))
    parsed.all_keywords   = list(set(parsed.skills))

    # ── Upgrade 4: Detect sparse/thin sections for onboarding hints ───────────
    sparse = []
    if not parsed.experience or all(len(e.bullets) < 2 for e in parsed.experience):
        sparse.append("experience")
    if len(parsed.projects) < 2:
        sparse.append("projects")
    if len(parsed.skills) < 5:
        sparse.append("skills")
    if not parsed.summary or len(parsed.summary) < 30:
        sparse.append("summary")
    parsed.sparse_sections = sparse

    logger.info(f"Parsed '{parsed.name}' | sections={parsed.sections_found} | sparse={sparse}")
    return parsed
